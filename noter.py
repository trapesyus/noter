#Özetimiz çıkarıldıktan sonra açmamız için gereken modülümüz
import os

#REGEX Kütüphanelerimiz
import re

#Image to Text Kütüphanelerimiz
from PIL import Image
import numpy as np
import cv2
import pytesseract

#Dosya Yolu Belirleme
from tkinter import Tk
from tkinter.filedialog import askopenfile
from tkinter.filedialog import asksaveasfile
from tkinter import ttk

#Arayüz Kütüphanemiz
import pyautogui

# PDF to TEXT
import fitz

# DOCX to TEXT
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches

# Resim Dosyasının Özeti
def imageToText():
    liste = list()

    pyautogui.alert(text='Lütfen Resim Dosyasını Seçin.', title='NOTER', button="Tamam")
    Tk().withdraw()
    yol = askopenfile()

    kayit_yol = yol.name.split(".jpg")
    docx.Document().save(kayit_yol[0] + "özet.docx")
    doc_ozet = docx.Document(kayit_yol[0] + "özet.docx")

    resim = cv2.imread(yol.name)

    screen_gray = cv2.cvtColor(resim, cv2.COLOR_BGR2GRAY)
    cv2.imwrite(kayit_yol[0] + "siyahbeyaz.jpg", screen_gray)

    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    tessdata_dir_config = "--tessdata-dir 'C:\Program Files\Tesseract-OCR\\tessdata'"



    metin = pytesseract.image_to_string(Image.open(kayit_yol[0] + "siyahbeyaz.jpg"), lang="tur", config=tessdata_dir_config)
    kmetin = metin.replace("\n", " ")

    os.remove(kayit_yol[0] + "siyahbeyaz.jpg")

    # Sayfa Düzeni
    sections = docx.Document().sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bot_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    #Kullanıcıdan istediği kadar anahtar kelime girmesini bu bölümde sağlıyoruz.
    while True:
        kelime = pyautogui.prompt(text='Lütfen Anahtar Kelimeyi Girin.', title='NOTER', default='')
        if kelime == "":
            break
        liste.append(kelime.lower())

    #Anahtar Kelimemiz ile Metinde Arama İşlemini Gerçekleştiriyoruz.
    for i in liste:
        notlar = re.findall(".*" + i + ".*" + "\n", kmetin.replace("." or "?" or "!", "\n").lower())
        if notlar: #Eğer Arama İşleminde Anahtar Kelimemiz Bulunmuşsa  Resmin özetini word dosyasına yazdırarak kaydediyoruz.

            doc_ozet.add_heading("\"{}\" Anahtar Kelimesi ile Oluşturulan Özet:\n".format(i.upper()), level=2)

            for j in notlar:

                paragraf = doc_ozet.add_paragraph(j.capitalize())
                if paragraf.style:
                    pass
                else:
                    paragraf.style = doc_ozet.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
                    font = paragraf.style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    font.color.rgb = RGBColor(0, 0, 0)
                doc_ozet.save(kayit_yol[0] + "özet.docx")

    open_folder = pyautogui.confirm(text='Dosyayı görüntülemek ister misiniz?', title='NOTER',
                                    buttons=["Evet", "Hayır"]);
    if open_folder == "Evet":
        os.startfile(kayit_yol[0] + "özet.docx")
    else:
        pass

# Metin Belgesinin Özeti
def textToText():
    liste = list()

    pyautogui.alert(text='Lütfen Metin Dosyasını Seçin.', title='NOTER', button="Tamam")
    Tk().withdraw()
    yol = askopenfile()

    dosya = open(yol.name, "r", encoding="utf-8")
    metin = dosya.read()
    dosya.close()

    kayit_yol = yol.name.split(".txt")
    dosya_ozet = open(kayit_yol[0] + "özet.txt", "w", encoding="utf-8")

    kmetin = metin.replace("\n", " ")


    #Kullanıcıdan istediği kadar anahtar kelime girmesini bu bölümde sağlıyoruz.
    while True:
        kelime = pyautogui.prompt(text='Lütfen Anahtar Kelimeyi Girin.', title='NOTER', default='')
        if kelime == "":
            break
        liste.append(kelime.lower())

    # Anahtar Kelimemiz ile Metinde Arama İşlemini Gerçekleştiriyoruz.
    for i in liste:
        notlar = re.findall(".*" + i + ".*" + "\n", metin.replace("." or "?" or "!", "\n").lower())
        # Eğer Arama İşleminde Anahtar Kelimemiz Bulunmuşsa  metin belgesinin özetini yeni bir metin belgesine yazdırarak kaydediyoruz.

        if notlar:
            dosya_ozet.write("\"{}\" Anahtar Kelimesi ile Oluşturulan Özet:\n".format(i.upper()))
            for j in notlar:
                dosya_ozet.write(j.capitalize())
    dosya_ozet.close()
    open_folder = pyautogui.confirm(text='Dosyayı görüntülemek ister misiniz?', title='NOTER',
                                    buttons=["Evet", "Hayır"])
    if open_folder == "Evet":
        os.startfile(kayit_yol[0] + "özet.txt")
    else:
        pass

#Word Dosyasının Özetini Çıkarma
def docxToDocx():
    kmetin = ""
    liste = list()

    pyautogui.alert(text='Lütfen Word Dosyasını Seçin.', title='NOTER', button="Tamam")
    Tk().withdraw()
    yol = askopenfile()

    doc = docx.Document(yol.name)
    kayit_yolu = yol.name.split(".docx")
    docx.Document().save(kayit_yolu[0] + "özet.docx")
    doc2 = docx.Document(kayit_yolu[0] + "özet.docx")
    metin = doc.paragraphs

    # Sayfa Düzeni
    sections = docx.Document().sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for i in metin:
        kmetin += i.text + " "

    # Kullanıcıdan istediği kadar anahtar kelime girmesini bu bölümde sağlıyoruz.
    while True:
        kelime = pyautogui.prompt(text='Lütfen Anahtar Kelimeyi Girin.', title='NOTER', default='')
        if kelime == "":
            break
        liste.append(kelime.lower())

    # Anahtar Kelimemiz ile Metinde Arama İşlemini Gerçekleştiriyoruz.
    for i in liste:
        notlar = re.findall(".*" + i + ".*" + "\n", kmetin.replace("." or "?" or "!", "\n".lower()))
        # Eğer Arama İşleminde Anahtar Kelimemiz Bulunmuşsa  word belgesinin özetini yeni bir word belgesine yazdırarak kaydediyoruz.

        if notlar:
            doc2.add_heading("\"{}\" Anahtar Kelimesi ile Oluşturulan Özet:".format(i.upper()), level=2)
            for j in notlar:
                paragraf = doc2.add_paragraph(j)

                if paragraf.style:
                    pass
                else:
                    paragraf.style = doc2.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
                    font = paragraf.style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    font.color.rgb = RGBColor(0, 0, 0)
                doc2.save(kayit_yolu[0] + "özet.docx")

    open_folder = pyautogui.confirm(text='Dosyayı görüntülemek ister misiniz?', title='NOTER',
                                    buttons=["Evet", "Hayır"]);
    if open_folder == "Evet":
        os.startfile(kayit_yolu[0] + "özet.docx")
    else:
        pass

#PDF Dosyasının Özetini Çıkarma
def pdfToPDF():
    metin = ""
    liste = list()

    pyautogui.alert(text='Lütfen PDF Dosyasını Seçin.', title='NOTER', button="Tamam")
    Tk().withdraw()
    yol = askopenfile()

    kayit_yol = yol.name.split(".pdf")
    docx.Document().save(kayit_yol[0] + "özet.docx")
    doc_ozet = docx.Document(kayit_yol[0] + "özet.docx")

    # Sayfa Düzeni
    sections = docx.Document().sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bot_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Kullanıcıdan istediği kadar anahtar kelime girmesini bu bölümde sağlıyoruz.
    while True:
        kelime = pyautogui.prompt(text='Lütfen Anahtar Kelimeyi Girin.', title='NOTER', default='')
        if kelime == "":
            break
        liste.append(kelime.lower())

    doc = fitz.open(yol.name)  # open document
    for page in doc:  # iterate the document pages #buradan devam et
        text = page.get_text().encode("utf8")  # get plain text (is in UTF-8)
        metin = metin + text.decode('utf-8')

    kmetin = metin.replace("\n", " ")

    # Anahtar Kelimemiz ile Metinde Arama İşlemini Gerçekleştiriyoruz.
    for i in liste:
        notlar = re.findall(".*" + i + ".*" + "\n", kmetin.replace("." or "?" or "!", "\n").lower())

        if notlar:
            doc_ozet.add_heading("\"{}\" Anahtar Kelimesi ile Oluşturulan Özet:\n".format(i.upper()), level=2)
            for j in notlar:
                paragraf = doc_ozet.add_paragraph(j.capitalize())
                if paragraf.style:
                    pass
                else:
                    paragraf.style = doc_ozet.styles.add_style('Style Name', WD_STYLE_TYPE.PARAGRAPH)
                    font = paragraf.style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    font.color.rgb = RGBColor(0, 0, 0)
                doc_ozet.save(kayit_yol[0] + "özet.docx")

    open_folder = pyautogui.confirm(text='Dosyayı görüntülemek ister misiniz?', title='NOTER',
                                        buttons=["Evet", "Hayır"]);
    if open_folder == "Evet":
            os.startfile(kayit_yol[0] + "özet.docx")
    else:
        pass

#Kullanıcıdan hangi özet çıkarma işlemini yapacağını aldığımız bölüm
def girisEkrani():
    islem = pyautogui.confirm(text='\tHoşgeldiniz\nLütfen Yapmak İstediğiniz İşlemi Seçin:', title='NOTER',
                              buttons=["Pdf Dosyası", "Word Dosyası", "Metin Belgesi", "Resim Dosyası"])
    if islem == "Pdf Dosyası":
        pdfToPDF()
    elif islem == "Word Dosyası":
        docxToDocx()
    elif islem == "Metin Belgesi":
        textToText()
    elif islem == "Resim Dosyası":
        imageToText()

#Programı girisEkrani() fonksiyonunu çağırarak başlatıyoruz
girisEkrani()
